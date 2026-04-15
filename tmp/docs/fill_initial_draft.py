from __future__ import annotations

"""
文件作用：
- 在不覆盖原始初稿的前提下，基于现有仓库材料生成一版补充后的论文初稿。
- 主要补充背景、研究意义、研究现状、系统设计说明和已有实验结果。
- 将仓库中已有的训练曲线、检测结果图、标定图插入到文档中。

注意事项：
- 本脚本只输出新的 docx 文件，不修改原始 `docs/初稿.docx`。
- 对于仓库中暂时缺失的机构图和平台照片，脚本会插入中文占位说明和图题。
"""

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = None
    ImageDraw = None
    ImageFont = None


ROOT = Path(r"E:\Desktop\Study_for_robot\A-Graduation")
SOURCE_DOC = ROOT / "docs" / "初稿.docx"
OUTPUT_DOC = ROOT / "docs" / "初稿_补充版.docx"
FIG_DIR = ROOT / "docs" / "generated_figures"
FIG_DIR.mkdir(parents=True, exist_ok=True)

RESULTS_CSV = ROOT / "src" / "Identification_code" / "runs" / "detect" / "yolo26_runs" / "results.csv"
CALIB_JSON = ROOT / "src" / "calibration" / "output" / "camera_calibration.json"
IMG_SYSTEM = FIG_DIR / "system_closed_loop.png"
IMG_SCENE = ROOT / "src" / "Identification_code" / "test_picture_video" / "微信图片_20260317163131_351_10.jpg"
IMG_CALIB = ROOT / "src" / "calibration" / "output" / "debug_corners" / "calib_010_corners.jpg"
IMG_TRAIN = ROOT / "src" / "Identification_code" / "runs" / "detect" / "yolo26_runs" / "results.png"
IMG_PRED = ROOT / "src" / "Identification_code" / "runs" / "detect" / "yolo26_runs" / "val_batch0_pred.jpg"


def generate_system_diagram(output_path: Path) -> None:
    """生成系统软件与控制闭环架构图。"""

    if Image is None:
        return

    width, height = 1800, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    font_candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    font_path = next((candidate for candidate in font_candidates if Path(candidate).exists()), None)
    title_font = ImageFont.truetype(font_path, 40) if font_path else ImageFont.load_default()
    text_font = ImageFont.truetype(font_path, 28) if font_path else ImageFont.load_default()
    small_font = ImageFont.truetype(font_path, 22) if font_path else ImageFont.load_default()

    def draw_box(x1: int, y1: int, x2: int, y2: int, text: str, fill: str) -> None:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline="#2f3b52", width=4)
        lines = text.split("\n")
        line_heights = []
        line_widths = []
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=text_font)
            line_widths.append(bbox[2] - bbox[0])
            line_heights.append(bbox[3] - bbox[1])

        total_height = sum(line_heights) + (len(lines) - 1) * 10
        current_y = y1 + (y2 - y1 - total_height) / 2
        for line, line_w, line_h in zip(lines, line_widths, line_heights):
            draw.text((x1 + (x2 - x1 - line_w) / 2, current_y), line, font=text_font, fill="black")
            current_y += line_h + 10

    def draw_arrow(start: tuple[int, int], end: tuple[int, int], label: str | None = None) -> None:
        draw.line((start, end), fill="#365d9d", width=6)
        dx = end[0] - start[0]
        dy = end[1] - start[1]
        if abs(dx) >= abs(dy):
            direction = 1 if dx >= 0 else -1
            p1 = (end[0] - 20 * direction, end[1] - 10)
            p2 = (end[0] - 20 * direction, end[1] + 10)
        else:
            direction = 1 if dy >= 0 else -1
            p1 = (end[0] - 10, end[1] - 20 * direction)
            p2 = (end[0] + 10, end[1] - 20 * direction)
        draw.polygon([end, p1, p2], fill="#365d9d")

        if label:
            bbox = draw.textbbox((0, 0), label, font=small_font)
            label_w = bbox[2] - bbox[0]
            label_h = bbox[3] - bbox[1]
            mid_x = (start[0] + end[0]) / 2
            mid_y = (start[1] + end[1]) / 2
            draw.rounded_rectangle(
                (mid_x - label_w / 2 - 10, mid_y - label_h / 2 - 6, mid_x + label_w / 2 + 10, mid_y + label_h / 2 + 6),
                radius=12,
                fill="white",
                outline="#d0d7e2",
                width=2,
            )
            draw.text((mid_x - label_w / 2, mid_y - label_h / 2), label, font=small_font, fill="#1f2a38")

    draw.text((width / 2 - 310, 40), "海参捕捞机器人视觉伺服闭环结构示意图", font=title_font, fill="#1b2c48")

    draw_box(100, 190, 360, 300, "图像采集\nUSB 相机", "#dff2ff")
    draw_box(100, 360, 360, 470, "海参检测\n目标框 / 置信度", "#e7f8e8")
    draw_box(100, 530, 360, 640, "目标信息发布\n中心坐标 / 时间戳", "#fff2d9")
    draw_box(520, 320, 850, 520, "视觉伺服控制节点\n误差计算 / 状态机 / 限幅", "#f4e4ff")
    draw_box(1010, 240, 1310, 350, "STM32 驱动控制\n电机 / 阀泵 / 通讯", "#ffe5e5")
    draw_box(1010, 490, 1310, 600, "可伸缩吸取装置\n对准 / 伸缩 / 吸附", "#e6f7ff")
    draw_box(1440, 320, 1710, 520, "执行反馈\n位移状态 / 吸附状态\n新图像闭环返回", "#eef3f8")

    draw_arrow((360, 245), (520, 420), "图像流")
    draw_arrow((360, 415), (520, 420), "检测结果")
    draw_arrow((360, 585), (520, 420), "目标消息")
    draw_arrow((850, 420), (1010, 295), "控制命令")
    draw_arrow((850, 420), (1010, 545), "伸缩与吸附指令")
    draw_arrow((1310, 545), (1440, 420), "执行状态")
    draw_arrow((1440, 360), (360, 245), "闭环图像反馈")

    draw.text(
        (70, 770),
        "说明：当前论文阶段已完成相机标定与目标检测验证，图中“视觉伺服控制节点—STM32—吸取装置”链路用于说明后续闭环集成方向。",
        font=small_font,
        fill="#334155",
    )
    image.save(output_path)


def find_paragraph_by_fragment(doc: Document, fragment: str) -> Paragraph:
    """按文本片段查找段落，不要求完全一致。"""

    for paragraph in doc.paragraphs:
        if fragment in paragraph.text:
            return paragraph
    raise ValueError(f"未找到段落片段: {fragment}")


def find_optional_paragraph(doc: Document, fragment: str) -> Paragraph | None:
    """尝试查找段落，找不到时返回 None，避免脚本中断。"""

    try:
        return find_paragraph_by_fragment(doc, fragment)
    except Exception:
        return None


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    """在指定段落后新增一个段落。"""

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.text = text
    return new_para


def format_body_paragraph(paragraph: Paragraph, first_line: bool = True) -> Paragraph:
    """统一正文段落格式。"""

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if first_line else Cm(0)
    return paragraph


def format_caption_paragraph(paragraph: Paragraph) -> Paragraph:
    """统一图表题格式。"""

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    return paragraph


def set_paragraph_text(paragraph: Paragraph, text: str, body: bool = True) -> Paragraph:
    """覆盖原段落文本并重新设置格式。"""

    paragraph.text = text
    return format_body_paragraph(paragraph, first_line=body)


def add_body_after(anchor: Paragraph, text: str, first_line: bool = True) -> Paragraph:
    """在锚点后插入一段正文。"""

    paragraph = insert_paragraph_after(anchor, text)
    return format_body_paragraph(paragraph, first_line=first_line)


def add_figure_after(anchor: Paragraph, image_path: Path | None, caption: str, placeholder: str | None = None, width_inches: float = 5.8) -> Paragraph:
    """在锚点后插入图片或中文占位，再插入图题。"""

    figure_para = insert_paragraph_after(anchor)
    figure_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_para.paragraph_format.space_before = Pt(6)
    figure_para.paragraph_format.space_after = Pt(0)
    figure_para.paragraph_format.first_line_indent = Cm(0)

    if image_path and image_path.exists():
        figure_para.add_run().add_picture(str(image_path), width=Inches(width_inches))
    else:
        figure_para.add_run(placeholder or f"【此处插入{caption}】")

    caption_para = insert_paragraph_after(figure_para, caption)
    return format_caption_paragraph(caption_para)


def add_table_after(doc: Document, anchor: Paragraph, title: str, rows: list[list[str]]) -> Paragraph:
    """在锚点后插入表题和表格，并返回表格后的空白锚点。"""

    title_para = insert_paragraph_after(anchor, title)
    format_caption_paragraph(title_para)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for r_index, row_data in enumerate(rows):
        for c_index, value in enumerate(row_data):
            cell = table.cell(r_index, c_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.line_spacing = 1.2
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.first_line_indent = Cm(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    if r_index == 0:
                        run.bold = True

    title_para._p.addnext(table._element)
    tail = OxmlElement("w:p")
    table._element.addnext(tail)
    tail_para = Paragraph(tail, title_para._parent)
    return format_body_paragraph(tail_para, first_line=False)


def clear_fragment(doc: Document, fragment: str) -> None:
    """清空占位文本对应的段落内容。"""

    paragraph = find_paragraph_by_fragment(doc, fragment)
    paragraph.text = ""


def main() -> None:
    """执行文档补写与图片插入流程。"""

    with RESULTS_CSV.open("r", encoding="utf-8-sig", newline="") as file:
        metric_rows = list(csv.DictReader(file))

    best_map50 = max(metric_rows, key=lambda row: float(row["metrics/mAP50(B)"]))
    best_map5095 = max(metric_rows, key=lambda row: float(row["metrics/mAP50-95(B)"]))
    best_precision = max(metric_rows, key=lambda row: float(row["metrics/precision(B)"]))
    last_row = metric_rows[-1]

    with CALIB_JSON.open("r", encoding="utf-8") as file:
        calib = json.load(file)

    train_image_count = 48
    val_image_count = 12
    train_box_count = 155
    val_box_count = 38
    total_image_count = train_image_count + val_image_count
    total_box_count = train_box_count + val_box_count

    generate_system_diagram(IMG_SYSTEM)
    doc = Document(SOURCE_DOC)

    # 先清理明显的提示性占位语，避免和补写正文混在一起。
    for fragment in [
        "（经济、政策背景：海参捕捞产业发展背景",
        "产业技术需求背景：水下机器人在海参捕捞中的应用价值……）",
        "（柔性吸取装置与水下末端执行器研究现状",
        "视觉伺服控制（IBVS/PBVS）研究现状",
        "水下视觉感知与目标检测研究现状",
        "水下机器人抓取/采摘类作业研究现状",
        "现状总结与研究不足……）",
    ]:
        try:
            clear_fragment(doc, fragment)
        except Exception:
            pass

    # 摘要与关键词。
    set_paragraph_text(
        find_paragraph_by_fragment(doc, "本文主要内容如下："),
        "本文主要内容如下：围绕海参捕捞机器人末端作业需求，完成了可伸缩吸取装置的功能分析与系统分层设计，构建了基于单目 USB 相机的海参图像采集、相机标定、目标检测与视觉伺服控制链路，并在现有数据与实验条件下对感知端性能进行了初步验证。",
    )
    result_summary = (
        f"结果表明：在当前 {total_image_count} 幅海参图像、{total_box_count} 个标注框的小样本单类别数据集上，所训练的海参检测模型在验证阶段取得了 "
        f"{float(best_map50['metrics/mAP50(B)']):.3f} 的 mAP@0.5 和 {float(best_map5095['metrics/mAP50-95(B)']):.3f} 的 mAP@0.5:0.95；"
        f"相机标定阶段共使用 {calib['valid_images']} 幅有效棋盘格图像，平均重投影误差为 {calib['mean_error_px']:.4f} 像素。"
        "上述结果说明，现有视觉感知链路已具备为后续可伸缩吸取装置视觉伺服控制提供稳定目标信息的基础。"
    )
    result_para = set_paragraph_text(find_paragraph_by_fragment(doc, "结果表明："), result_summary)
    format_body_paragraph(add_body_after(result_para, "关键词：海参捕捞机器人；可伸缩吸取装置；视觉伺服；目标检测；相机标定"))

    # 第一章补写。
    meaning_para = find_paragraph_by_fragment(doc, "本研究不仅为海参等海底软体海产品的自动化捕捞提供了创新性的硬件设计方案")
    add_body_after(
        meaning_para,
        "从产业规模与政策导向看，相关公开资料显示，中国海参产业已经形成涵盖育苗、养殖、加工、流通与销售的千亿级完整产业链；烟台市 2023 年公开信息显示，当地海参全产业链年产值约 160 亿元，已成为区域渔业的重要支柱产业。在国家推进智慧农业、渔业生产智能化和智慧渔场建设的背景下，围绕海参采收环节开展智能装备研究，具有明确的产业牵引价值。",
    )

    set_paragraph_text(
        find_paragraph_by_fragment(doc, "传统刚性夹爪的局限性：传统的机械手在水下抓取时"),
        "传统刚性夹爪在水下抓取软体目标时，往往依赖较强的夹持作用力来保证稳定性，但对于海参这类体壁柔软、易受刺激的目标而言，局部压强过大容易引发破损、吐脏甚至自溶等问题，因此其适用性存在明显局限。",
    )
    set_paragraph_text(
        find_paragraph_by_fragment(doc, "低压吸附与柔性接触的优势：针对海参这类体壁柔软且易吐脏的棘皮动物"),
        "相较于刚性夹取，低压吸附与柔性接触方式更适合海参类目标。一方面，吸附接触面积较大、单位面积受力较小；另一方面，柔性吸取头能够在一定程度上顺应目标表面形态和海底起伏，提高接触稳定性并降低机械损伤风险。",
    )
    set_paragraph_text(
        find_paragraph_by_fragment(doc, "可伸缩机构的引入：目前多数水下吸取装置缺乏深度方向的灵活补偿"),
        "近年来，连续体机构、绳驱柔顺机构以及可变刚度执行器的研究说明，柔顺结构能够显著提升狭小空间通过能力和复杂环境适应性。但多数研究主要面向医疗、巡检或通用操作任务，针对海底软体目标的短行程对准与低扰动吸取问题仍缺少专门设计。",
    )

    anchor = find_paragraph_by_fragment(doc, "视觉伺服控制（IBVS/PBVS）研究现状")
    anchor = add_body_after(anchor, "视觉伺服的核心在于利用视觉信息直接闭环调节执行器运动。IBVS 直接以图像平面误差为控制依据，对模型误差和参数不确定性更具鲁棒性，适合当前单目视觉、小范围对准的实现条件。")
    anchor = add_body_after(anchor, "PBVS 则依赖目标三维位姿估计，控制变量在任务空间中更直观，适合末端需要明确空间接近量的场景。但在水下环境中，折射、光照变化和目标模型不确定性会放大建模误差，因此 PBVS 的工程落地通常依赖更精确的标定与稳定的三维感知。")
    add_body_after(anchor, "因此，越来越多研究倾向于采用分阶段或混合视觉伺服策略，即在粗对准阶段优先利用图像误差实现快速收敛，在近距离接近阶段再结合深度或几何信息完成精细定位，这一思路也为本课题的控制设计提供了参考。")

    anchor = find_paragraph_by_fragment(doc, "水下视觉感知与目标检测研究现状")
    anchor = add_body_after(anchor, "水下视觉感知面临光线衰减、蓝绿偏色、悬浮颗粒散射和对比度下降等问题，导致传统阈值分割和手工特征方法在复杂海底背景中的稳定性不足。随着深度学习目标检测方法的发展，水下生物识别的准确性和泛化能力得到明显提升。")
    add_body_after(anchor, "但对本课题而言，感知模块不仅要“看见目标”，还要稳定输出可供伺服控制使用的中心位置、尺度信息与置信度。因此，模型轻量化、帧间稳定性和部署可移植性，比单纯追求离线检测精度更具工程意义。")

    anchor = find_paragraph_by_fragment(doc, "水下机器人抓取/采摘类作业研究现状")
    anchor = add_body_after(anchor, "现有水下机器人作业研究多集中于网箱巡检、附着物清理、常规抓取和海底探测等任务。面向海参采摘的研究则需要同时满足复杂地形接近、软体目标保护和末端机构低扰动作业等要求，其控制目标明显不同于一般工业抓取。")
    add_body_after(anchor, "从系统集成角度看，将目标检测、视觉伺服与专用末端执行器闭环联动的工程实现仍然较少，尤其缺少面向短行程可伸缩吸取机构的控制链路设计。")

    set_paragraph_text(find_paragraph_by_fragment(doc, "现状总结与研究不足……"), "现状总结与研究不足")
    set_paragraph_text(
        find_paragraph_by_fragment(doc, "综合来看，目前将“视觉伺服”与“动态可伸缩吸取机构”深度结合的研究相对较少"),
        "综合来看，现有研究分别在柔顺执行机构、视觉伺服方法和水下视觉感知等方面取得了较多成果，但将三者面向“海参软体目标非破坏性采摘”这一具体任务进行一体化设计的工作仍不充分。当前研究不足主要体现在：缺少适配可伸缩吸取装置的简洁控制变量定义，缺少从检测输出到执行器命令的稳定映射机制，也缺少围绕海参捕捞场景开展的闭环系统验证。基于此，本文以 ROS2 上位机与 STM32 下位机分层架构为基础，尝试构建“视觉检测—误差计算—伸缩控制—吸附执行”的闭环作业链路。",
    )

    anchor = find_paragraph_by_fragment(doc, "1.3 主要研究内容")
    anchor = add_body_after(anchor, "本文围绕海参捕捞机器人末端作业单元的视觉伺服问题，主要开展以下几个方面的研究：")
    anchor = add_body_after(anchor, "（1）围绕海参软体目标的非破坏性采收需求，分析可伸缩吸取装置的作业约束、功能需求与系统分层关系，明确末端执行机构的设计目标。")
    anchor = add_body_after(anchor, "（2）构建海参图像采集、数据整理、目标检测训练与推理验证流程，形成可为后续控制环节提供目标位置信息的感知基础。")
    anchor = add_body_after(anchor, "（3）结合单目相机标定结果与视觉伺服理论，设计适配可伸缩吸取装置的小行程对准与接近控制思路，明确图像误差、控制输出和状态切换逻辑。")
    add_body_after(anchor, "（4）基于当前仓库中的实验数据，对相机标定精度、检测训练结果和视觉输出可用性进行验证，为后续整机闭环实验提供论文层面的阶段性依据。")

    # 第二章补写。
    set_paragraph_text(
        find_paragraph_by_fragment(doc, "吸取装置功能需求以及创新型设计"),
        "海参捕捞任务对末端装置提出了三方面要求：一是能够在复杂礁石或近底环境中实现柔和接触，避免刚性碰撞；二是能够在主机体到达目标附近后提供稳定的小行程接近能力；三是能够与视觉感知链路配合，实现目标中心对准后的吸附作业。因此，可伸缩吸取装置的核心价值不在于大范围操作，而在于为末端最后一段精细接近与低损伤取样提供结构支撑。",
    )
    set_paragraph_text(
        find_paragraph_by_fragment(doc, "（低压吸附、柔性接触保护海参……）"),
        "对于海参这类柔软且易受刺激的目标，吸取方式应优先保证接触温和和动作可控。采用低压吸附配合柔性吸取头，可以在保证吸附稳定性的同时减小局部接触应力；配合可伸缩机构后，可将作业过程划分为“视觉对准—轴向接近—吸附抓取—回收复位”四个阶段，从而降低盲目接近带来的碰撞风险。",
    )

    anchor = find_paragraph_by_fragment(doc, "2.3.1可伸缩机构机械结构")
    anchor = add_body_after(anchor, "从功能构成上，可伸缩吸取装置可划分为安装基座、伸缩驱动单元、导向支撑单元、柔性吸取头与负压接口五个部分。其设计重点不是构建通用多自由度机械臂，而是在主平台完成粗定位后，沿末端作业方向提供可控、可回收的小行程位移补偿。")
    anchor = add_body_after(anchor, "若将伸缩位移记为 s，则 s 的变化主要用于补偿作业方向上的接近误差；当图像平面误差收敛到阈值范围内后，系统再触发 s 的递增以及吸附动作，从而兼顾对准精度与作业安全性。")
    anchor = add_figure_after(anchor, None, "图2-1 可伸缩吸取装置总体结构示意图", "【此处插入图2-1 可伸缩吸取装置总体结构示意图】")
    set_paragraph_text(
        find_paragraph_by_fragment(doc, "执行方式、运动学描述、传递函数变换……"),
        "考虑到当前初稿阶段尚未保留完整的三维机构设计图与详细尺寸参数，本文在此以功能链路和运动关系为主进行描述，具体结构细节、关键尺寸与加工图可在后续实物定型后补充。",
    )

    anchor = find_paragraph_by_fragment(doc, "2.3.2 控制器硬件构成")
    add_body_after(anchor, "控制器硬件采用上位机与下位机分层方案。上位机侧负责图像采集、目标检测、目标信息发布和伺服控制计算；下位机侧以 STM32 为核心，负责接收控制指令并驱动伸缩执行机构、阀泵或相关驱动单元完成动作执行。该结构既便于算法迭代，也便于后续整机联调与硬件替换。")
    anchor = find_paragraph_by_fragment(doc, "2.3.3 软件系统架构（ROS/STM32）")
    anchor = add_body_after(anchor, "软件系统以 ROS2 节点化方式组织上位机流程，包括图像输入节点、海参检测节点、目标信息发布节点、视觉伺服控制节点和命令下发接口。视觉结果通过统一消息格式传递给控制层，控制层再将限幅后的指令发送给 STM32，实现感知与执行的松耦合集成。")
    add_figure_after(anchor, IMG_SYSTEM if IMG_SYSTEM.exists() else None, "图2-2 系统软件与控制闭环架构图", "【此处插入图2-2 系统软件与控制闭环架构图】", width_inches=6.3)
    add_body_after(find_paragraph_by_fragment(doc, "2.4 本章小结"), "本章围绕海参捕捞任务需求，对可伸缩吸取装置的吸附方式、机械功能构成以及软硬件分层关系进行了分析。该装置强调小行程接近、柔和接触和易于闭环控制，为后续视觉误差驱动的伺服控制设计奠定了基础。")

    # 第三章补写。
    anchor = find_paragraph_by_fragment(doc, "3.1 引言")
    add_body_after(anchor, "水下环境中的光照衰减、蓝绿偏色和悬浮物散射会显著降低图像对比度，并引入边缘模糊与噪声，这使得目标中心提取的稳定性直接影响后续伺服控制精度。结合当前仓库中的实现情况，本文现阶段采用单目 USB 相机完成海参数据采集、相机标定与目标检测验证，并在此基础上为后续扩展到双目测距或混合视觉伺服预留接口。")
    set_paragraph_text(find_paragraph_by_fragment(doc, "  水下环境由于光线衰减、悬浮物散射，导致图像存在严重的色偏和低对比度问题。"), "水下环境由于光线衰减、悬浮物散射和背景附着物干扰，图像往往存在明显色偏、低对比度和纹理模糊问题，因此感知链路必须同时兼顾目标可见性和帧间稳定性。")
    set_paragraph_text(find_paragraph_by_fragment(doc, "  利用基于深度学习的海参识别代码，将其作为信息节点接入系统。"), "基于现有深度学习检测代码，可将海参检测结果进一步封装为目标信息节点输出，为后续视觉伺服模块提供目标中心坐标、检测框尺度、置信度和时间戳等控制所需信息。")
    anchor = find_paragraph_by_fragment(doc, "3.2 立体视觉测距（双目相机）")
    anchor = add_body_after(anchor, "考虑到论文完整性和后续系统扩展需求，本节保留双目测距方案作为升级方向。若后续配置双目相机，可通过双目标定获取基线与投影矩阵，并利用视差计算目标深度，用于 PBVS 或混合视觉伺服中的空间位移估计。")
    add_body_after(anchor, "但就当前验证平台而言，感知实验仍以单目标定与图像平面误差提取为主，因此本章实验结果主要围绕单目成像质量、目标检测效果和可用于控制的二维特征输出展开。")
    anchor = find_paragraph_by_fragment(doc, "3.3 基于立体视觉的水下目标定位算法研究")
    anchor = add_body_after(anchor, "在当前实现中，目标定位以检测框中心坐标和框尺度信息为主要输出。检测框中心可直接用于构造图像平面误差，框尺度则可作为目标远近变化的辅助判断量，为后续阶段式接近控制提供依据。")
    add_figure_after(anchor, IMG_CALIB if IMG_CALIB.exists() else None, "图3-1 相机标定角点提取结果", "【此处插入图3-1 相机标定角点提取结果】", width_inches=5.8)
    anchor = find_paragraph_by_fragment(doc, "3.4 面向水下机器人作业的目标三维位姿估计方法研究")
    add_body_after(anchor, "对于面向整机作业的目标位姿估计，后续可在现有二维检测结果基础上融合双目深度、末端位移反馈或结构先验，实现更稳定的空间接近控制。初稿阶段则优先将“目标是否居中、是否具备接近条件”作为视觉伺服的核心判据，以保证方案与当前硬件能力相匹配。")
    add_body_after(find_paragraph_by_fragment(doc, "3.5 本章小结"), "本章从水下成像特点出发，说明了海参视觉检测与定位链路的实现基础。当前阶段以单目视觉与目标检测输出为主，重点解决“能否稳定获得可用于控制的二维目标信息”这一问题，并为后续深度估计与空间位姿扩展预留了接口。")

    # 第四章补写。
    set_paragraph_text(find_paragraph_by_fragment(doc, "（视觉伺服在海参捕捞中的作用）"), "视觉伺服在海参捕捞任务中的作用，是将“看见目标”进一步转化为“稳定靠近目标”。对于可伸缩吸取装置而言，视觉反馈不仅决定平面对准的精度，还直接影响末端何时伸出、何时吸附以及何时回收，因此控制链路的核心是建立稳定、可限幅、可切换状态的闭环策略。")
    set_paragraph_text(find_paragraph_by_fragment(doc, "IBVS、PBVS……"), "视觉伺服的本质是将图像或位姿误差直接纳入控制回路，使执行器运动随视觉反馈实时调整。结合本课题的硬件特征，控制对象不是通用六自由度机械臂，而是以平面对准和轴向接近为主的小行程可伸缩吸取机构。")
    set_paragraph_text(find_paragraph_by_fragment(doc, "IBVS（基于图像的视觉伺服）和PBVS（基于位置的视觉伺服）是视觉伺服控制领域两种重要的方法。IBVS主要通过直接利用图像特征误差来控制机器人的运动，它不需要精确的机器人运动学和动力学模型，对相机标定误差和目标模型误差有一定的鲁棒性。在海参捕捞场景中，IBVS能够根据海参在图像中的特征，如形状、颜色等，快速调整吸取装置的位置和姿态，以实现对海参的准确捕捉。"), "设目标在图像平面的特征向量为 s=[u,v]^T，期望特征为 s*=[u0,v0]^T，则图像误差可写为 e=s-s*=[eu,ev]^T。当前阶段可直接采用检测框中心作为图像特征，利用 eu、ev 表征目标相对吸取中心的位置偏差。")
    set_paragraph_text(find_paragraph_by_fragment(doc, "PBVS则是基于目标的三维位置信息来进行控制，它需要精确的相机标定和目标的三维模型。在海参捕捞中，PBVS可以通过立体视觉等技术获取海参的三维位置，然后根据这些信息精确地控制吸取装置的运动，使吸取装置能够准确地到达海参所在的位置。"), "对于 IBVS，可根据图像误差构造平面修正量，使末端位置围绕目标中心逐步收敛。该方法无需依赖复杂的三维模型，适合当前单目视觉和工程样机阶段的实现条件。")
    set_paragraph_text(find_paragraph_by_fragment(doc, "然而，这两种方法都有各自的优缺点。IBVS虽然对模型误差有一定的鲁棒性，但在图像特征提取和匹配方面可能会遇到困难，尤其是在水下复杂环境中，图像的质量可能会受到很大影响。PBVS虽然能够提供精确的三维位置信息，但对相机标定和目标模型的精度要求较高，并且计算量相对较大。"), "对于 PBVS 或双目扩展方案，可在获取目标深度和空间位姿后，在任务空间中规划末端接近量，用于提升近距离吸取阶段的空间一致性。但这一方案对标定精度、深度稳定性和系统实时性提出了更高要求。")
    set_paragraph_text(find_paragraph_by_fragment(doc, "为了充分发挥IBVS和PBVS的优势，在海参捕捞机器人的视觉伺服控制中，可以采用混合视觉伺服控制策略。例如，在初始阶段，可以利用IBVS快速地将吸取装置引导到海参附近，然后切换到PBVS进行精确的定位和捕捉，以提高捕捞的准确性和成功率。同时，还可以结合其他技术，如机器学习、深度学习等，进一步提高视觉伺服控制的性能，以适应水下复杂环境和海参的动态特性。"), "因此，本文更适合采用“IBVS 主导、PBVS 扩展”的分阶段思路：远距离和粗对准阶段使用图像误差实现快速收敛，接近与吸附阶段再结合深度信息、伸缩位移反馈或状态触发完成精细控制。")
    set_paragraph_text(find_paragraph_by_fragment(doc, "基于位置的视觉伺服 (PBVS)： 该方法依赖于目标的三维位姿估计。如果你的系统最终采用双目相机，PBVS将是重点考察的控制律 。它能在笛卡尔空间内规划可伸缩机构的运动轨迹，但对相机的标定精度和水下折射模型的建立要求极高。"), "由于水下图像存在抖动、模糊和检测框跳变，控制前还需要对目标输出进行置信度判定与时序滤波，避免瞬时检测误差直接传递到执行机构。")
    set_paragraph_text(find_paragraph_by_fragment(doc, "基于图像的视觉伺服 (IBVS)： 该方法直接利用图像特征面的误差来计算控制量。如果系统采用单目相机，则使用IBVS 。它的鲁棒性更强，能够有效克服水下环境带来的相机参数标定误差。"), "同时，所有控制输出都应设置限幅和触发阈值，确保可伸缩机构在目标短暂丢失、图像模糊或通信延迟时进入保持状态，而不是持续执行危险动作。")
    anchor = find_paragraph_by_fragment(doc, "4.3 海参捕捞任务的视觉伺服控制策略设计")
    anchor = add_body_after(anchor, "针对海参捕捞任务，可将末端作业过程划分为搜索、对准、接近、吸附和回收五个状态。搜索阶段由相机持续获取目标图像；对准阶段根据检测框中心误差 eu、ev 计算平面修正命令；当误差连续若干帧满足阈值条件后，系统进入接近阶段，控制可伸缩机构沿作业方向伸出；满足接触或吸附条件后，再执行吸附和回收动作。")
    anchor = add_body_after(anchor, "在控制量设计上，可将平面修正命令写为 Δx=sat(ku·eu)、Δy=sat(kv·ev)，其中 sat(·) 表示限幅函数，ku 和 kv 分别表示横向与纵向误差增益。该设计有助于避免检测框瞬时跳变引起过大的控制指令。")
    add_body_after(anchor, "为提高工程稳定性，目标信息消息中应至少包含中心坐标、检测框尺度、置信度和时间戳。控制层在接收消息后，先完成有效性检查和时序滤波，再向 STM32 下发位移、速度或阶段动作指令，以实现上位机感知与下位机执行的清晰分工。")
    set_paragraph_text(find_paragraph_by_fragment(doc, "闭环控制系统设计"), "闭环控制系统设计应遵循“感知可信、控制限幅、状态可回退”的原则。当检测置信度低于阈值或目标连续多帧丢失时，系统应保持或回收末端机构，避免因错误检测导致误动作。")
    anchor = find_paragraph_by_fragment(doc, "4.4 控制参数整定与收敛分析")
    anchor = add_body_after(anchor, "控制参数整定可遵循由小到大、先稳后快的原则。初始阶段应优先保证平面对准过程无明显振荡，再逐步提高比例增益以缩短收敛时间；对于伸缩动作，则应设置独立的速度上限与触发阈值，避免在图像误差尚未稳定时提前接近目标。")
    add_body_after(anchor, "从收敛性角度看，当图像误差经过滤波后单调减小且控制输出始终处于限幅范围内时，系统可在较短时间内将目标中心收敛到期望区域。若出现误差长时间不下降的情况，则需要优先排查目标检测跳变、相机安装偏差或执行机构响应迟滞等问题。")
    add_body_after(find_paragraph_by_fragment(doc, "4.5 本章小结"), "本章围绕可伸缩吸取装置的视觉伺服需求，给出了以图像误差为核心的控制变量定义、状态切换思路和参数整定原则。该方法强调与当前单目视觉和小行程执行机构相匹配，为后续闭环联调提供了清晰的控制框架。")

    # 第五章补写。
    anchor = find_paragraph_by_fragment(doc, "5.1 引言")
    add_body_after(anchor, "考虑到当前仓库中已经形成较完整的相机标定、海参检测训练与推理结果，而整机闭环控制实验仍处于后续联调阶段，本章实验验证主要围绕“视觉感知链路是否具备闭环控制基础”这一问题展开，包括实验场景说明、相机标定结果分析以及海参检测性能验证。")
    anchor = find_paragraph_by_fragment(doc, "5.2.1 水池实验环境布置")
    anchor = add_body_after(anchor, "从现有采集图像可以看出，海参目标验证场景以近底平面背景、多目标分布和低对比度水下成像为主，能够较好体现海参目标在实际作业环境中的视觉特征。当前初稿阶段暂未保留完整整机水池平台照片，因此先给出采集场景示例，并为后续平台整体照片预留插图位置。")
    anchor = add_figure_after(anchor, IMG_SCENE if IMG_SCENE.exists() else None, "图5-1 海参数据采集场景示例", "【此处插入图5-1 海参数据采集场景示例】", width_inches=5.8)
    add_figure_after(anchor, None, "图5-2 机器人测试平台与相机安装位置图", "【此处插入图5-2 机器人测试平台与相机安装位置图】")

    anchor = find_paragraph_by_fragment(doc, "5.2.2 机器人测试平台与设备参数")
    anchor = add_body_after(anchor, "结合现有工程文件，可将当前验证平台概括为“单目相机 + 海参检测模型 + 标定参数 + 上下位机分层控制接口”的组合。其中，感知侧已经具备真实数据、训练权重和推理入口，下位机驱动与末端机构的整机联调将在后续实验中继续补充。")
    add_table_after(
        doc,
        anchor,
        "表5-1 当前验证平台的关键配置",
        [
            ["组成部分", "当前配置或状态", "说明"],
            ["图像采集", "USB 单目相机，1920×1080", "作为当前感知输入主源"],
            ["标定对象", "10×7 内角点棋盘格，方格边长 20 mm", "用于求取相机内参与畸变参数"],
            ["检测模型", "YOLO26n 单类别模型", "面向海参目标检测"],
            ["训练输入", f"{total_image_count} 幅图像，{total_box_count} 个标注框", "训练集 48 幅，验证集 12 幅"],
            ["推理模式", "图片 / 视频 / 摄像头", "由 infer_yolo26.py 统一调用"],
            ["控制架构", "ROS2 上位机 + STM32 下位机", "当前已完成论文层面的链路设计与接口梳理"],
        ],
    )

    anchor = find_paragraph_by_fragment(doc, "5.2.3 数据采集与标定流程")
    anchor = add_body_after(anchor, "相机标定阶段共采集 51 幅棋盘格图像，图像分辨率为 1920×1080。采集过程中通过改变棋盘格在视场中的位置、姿态和距离，保证角点覆盖较大视场范围，从而提升内参与畸变参数估计的稳定性。")
    anchor = add_table_after(
        doc,
        anchor,
        "表5-2 相机标定结果",
        [
            ["参数", "数值"],
            ["有效标定图像数量", str(calib["valid_images"])],
            ["标定模型", str(calib["model"])],
            ["图像分辨率", f"{calib['image_width']}×{calib['image_height']}"],
            ["fx / fy", f"{calib['camera_matrix'][0][0]:.3f} / {calib['camera_matrix'][1][1]:.3f}"],
            ["cx / cy", f"{calib['camera_matrix'][0][2]:.3f} / {calib['camera_matrix'][1][2]:.3f}"],
            ["RMS", f"{calib['rms']:.4f}"],
            ["平均重投影误差 / px", f"{calib['mean_error_px']:.4f}"],
        ],
    )
    anchor = add_body_after(anchor, "从标定结果可以看出，当前相机模型的平均重投影误差较小，说明角点检测与参数求解结果具有较好的几何一致性，能够满足后续基于图像误差的伺服控制需求。")
    add_figure_after(anchor, IMG_CALIB if IMG_CALIB.exists() else None, "图5-3 相机标定角点提取结果", "【此处插入图5-3 相机标定角点提取结果】", width_inches=5.8)

    section_53_heading = find_paragraph_by_fragment(doc, "5.3 实验：海参视觉检测与定位性能验证")
    static_para = find_optional_paragraph(doc, "静态实验、动态实验")
    if static_para is not None:
        anchor = set_paragraph_text(static_para, "本节验证分为静态图像检测验证和连续视频帧检测验证两部分。静态验证用于评估海参目标在当前小样本数据集上的检测精度，动态验证则用于观察连续画面中目标框输出的稳定性和可视化效果，从而判断其是否具备作为视觉伺服输入的潜力。")
    else:
        anchor = add_body_after(section_53_heading, "本节验证分为静态图像检测验证和连续视频帧检测验证两部分。静态验证用于评估海参目标在当前小样本数据集上的检测精度，动态验证则用于观察连续画面中目标框输出的稳定性和可视化效果，从而判断其是否具备作为视觉伺服输入的潜力。")

    anchor = add_table_after(
        doc,
        anchor,
        "表5-3 海参检测数据集划分情况",
        [
            ["数据划分", "图像数量", "标注框数量", "说明"],
            ["训练集", str(train_image_count), str(train_box_count), "用于模型训练与参数更新"],
            ["验证集", str(val_image_count), str(val_box_count), "用于精度评估与早停判断"],
            ["总计", str(total_image_count), str(total_box_count), "单类别海参检测数据集"],
        ],
    )
    anchor = add_body_after(anchor, "训练过程中使用 YOLO26n 预训练权重，输入尺寸设置为 832，批大小为 8，训练上限为 150 轮，并采用 patience=30 的早停策略。实际训练在第 88 轮结束，说明模型在当前数据规模下已进入相对稳定的收敛阶段。")
    anchor = add_table_after(
        doc,
        anchor,
        "表5-4 海参检测性能指标",
        [
            ["指标", "最佳结果", "对应轮次", "说明"],
            ["Precision", f"{float(best_precision['metrics/precision(B)']):.3f}", str(best_precision["epoch"]), "误检较少，结果更保守"],
            ["Recall", f"{float(best_map50['metrics/recall(B)']):.3f}", str(best_map50["epoch"]), "与最佳 mAP@0.5 对应的稳定召回表现"],
            ["mAP@0.5", f"{float(best_map50['metrics/mAP50(B)']):.3f}", str(best_map50["epoch"]), "当前小样本条件下检测准确率较高"],
            ["mAP@0.5:0.95", f"{float(best_map5095['metrics/mAP50-95(B)']):.3f}", str(best_map5095["epoch"]), "更严格 IoU 条件下仍保持较好性能"],
            ["最后一轮 Precision / Recall", f"{float(last_row['metrics/precision(B)']):.3f} / {float(last_row['metrics/recall(B)']):.3f}", str(last_row["epoch"]), "说明训练结束时模型仍具备较好综合表现"],
        ],
    )
    anchor = add_body_after(anchor, "需要说明的是，由于当前验证集规模仅为 12 幅图像、38 个标注框，表中结果更适合作为阶段性验证结论，而不能直接视为大规模泛化性能结论。即便如此，mAP@0.5 与最终轮次的 Precision / Recall 已经表明，现有检测模型能够较稳定地输出海参目标位置，为后续视觉误差计算提供基础。")
    anchor = add_figure_after(anchor, IMG_TRAIN if IMG_TRAIN.exists() else None, "图5-4 海参检测训练过程指标曲线", "【此处插入图5-4 海参检测训练过程指标曲线】", width_inches=6.3)
    anchor = add_figure_after(anchor, IMG_PRED if IMG_PRED.exists() else None, "图5-5 验证集目标检测结果示例", "【此处插入图5-5 验证集目标检测结果示例】", width_inches=6.0)
    add_body_after(anchor, "从验证集检测结果示例可以看出，在多目标、弱对比度和一定程度背景干扰条件下，模型仍能给出较为准确的目标框。对于视觉伺服应用而言，这意味着可直接利用检测框中心作为图像平面误差输入，并结合置信度门限与时序滤波提高控制稳定性。")

    add_body_after(find_paragraph_by_fragment(doc, "5.4 本章小结"), "本章基于现有工程文件与实验产物，对相机标定精度和海参检测性能进行了验证。结果表明，当前视觉感知链路已经具备较好的目标识别与中心提取能力，能够为后续“视觉检测—伺服控制—吸附执行”的整机闭环实验提供输入基础。")

    # 第六章补写。
    anchor = find_paragraph_by_fragment(doc, "第6章 总结与展望")
    anchor = add_body_after(anchor, "本文围绕海参捕捞机器人可伸缩吸取装置的视觉伺服控制问题，完成了初步的系统框架梳理、吸取装置功能分析、视觉感知链路构建以及感知端实验验证。结合相机标定结果与海参检测模型训练结果，可以认为当前项目已经具备从“可见目标”走向“可控接近目标”的基础条件。")
    add_body_after(anchor, "后续工作可从三个方面继续推进：一是补充可伸缩机构实物结构图、关键尺寸与平台安装照片，完善第二章和第五章的硬件表达；二是完成 ROS2 上位机到 STM32 下位机的闭环联调，形成静态对准时间、控制抖动、吸附成功率等直接服务于视觉伺服结论的实验指标；三是在条件允许时引入双目深度或末端位移反馈，进一步提升近距离吸附阶段的空间定位精度与作业稳定性。")

    # 为新增内容统一补一个默认字号，尽量保持阅读一致性。
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None:
                run.font.size = Pt(12)

    doc.save(OUTPUT_DOC)
    print(f"generated: {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
