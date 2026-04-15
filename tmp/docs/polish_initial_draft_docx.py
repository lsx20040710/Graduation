from __future__ import annotations

"""
文件作用：
- 对 `初稿_补充版.docx` 做第二轮论文格式润色。
- 重点处理公式源码感、字体颜色不统一、表格无边框等问题。
- 输出新的润色版文档，保留前一版作为备份。
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"E:\Desktop\Study_for_robot\A-Graduation")
SOURCE_DOC = ROOT / "docs" / "初稿_补充版.docx"
OUTPUT_DOC = ROOT / "docs" / "初稿_补充版_润色.docx"


def set_run_font(run, chinese_font: str, ascii_font: str, size_pt: float, bold: bool = False) -> None:
    """统一设置一个文本片段的中英文字体、字号和颜色。"""

    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese_font)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_fonts(paragraph, chinese_font: str, ascii_font: str, size_pt: float, bold: bool = False) -> None:
    """统一设置一个段落内所有文本片段的字体样式。"""

    for run in paragraph.runs:
        set_run_font(run, chinese_font, ascii_font, size_pt, bold=bold)


def is_chapter_heading(text: str) -> bool:
    """判断是否为章标题。"""

    return bool(re.match(r"^第\d+章", text))


def is_section_heading(text: str) -> bool:
    """判断是否为节或小节标题。"""

    return bool(re.match(r"^\d+\.\d+(\.\d+)?", text))


def is_caption(text: str) -> bool:
    """判断是否为图题或表题。"""

    return bool(re.match(r"^[图表]\d+-\d+", text))


def replace_formula_like_paragraphs(doc: Document) -> None:
    """把明显像源码的公式文本改成更适合正文的表达。"""

    replacements = {
        "设目标在图像平面的特征向量为 s=[u,v]^T，期望特征为 s*=[u0,v0]^T，则图像误差可写为 e=s-s*=[eu,ev]^T。当前阶段可直接采用检测框中心作为图像特征，利用 eu、ev 表征目标相对吸取中心的位置偏差。":
        "设目标中心当前像素坐标为 (u, v)，期望中心坐标为 (u₀, v₀)，则横向误差和纵向误差可分别写为 e_u = u - u₀、e_v = v - v₀。当前阶段可直接采用检测框中心作为图像特征，并据此构造视觉伺服控制误差。",
        "针对海参捕捞任务，可将末端作业过程划分为搜索、对准、接近、吸附和回收五个状态。搜索阶段由相机持续获取目标图像；对准阶段根据检测框中心误差 eu、ev 计算平面修正命令；当误差连续若干帧满足阈值条件后，系统进入接近阶段，控制可伸缩机构沿作业方向伸出；满足接触或吸附条件后，再执行吸附和回收动作。":
        "针对海参捕捞任务，可将末端作业过程划分为搜索、对准、接近、吸附和回收五个状态。搜索阶段由相机持续获取目标图像；对准阶段根据检测框中心误差 e_u、e_v 计算平面修正命令；当误差连续若干帧满足阈值条件后，系统进入接近阶段，控制可伸缩机构沿作业方向伸出；满足接触或吸附条件后，再执行吸附和回收动作。",
        "在控制量设计上，可将平面修正命令写为 Δx=sat(ku·eu)、Δy=sat(kv·ev)，其中 sat(·) 表示限幅函数，ku 和 kv 分别表示横向与纵向误差增益。该设计有助于避免检测框瞬时跳变引起过大的控制指令。":
        "在控制量设计上，平面修正量采用“比例计算加限幅约束”的方式生成：横向修正量由横向误差经比例增益处理后得到，纵向修正量由纵向误差经比例增益处理后得到，再统一通过限幅函数约束输出幅值。其中，k_u 和 k_v 分别表示横向与纵向比例增益。该设计有助于避免检测框瞬时跳变引起过大的控制指令。",
        "需要说明的是，由于当前验证集规模仅为 12 幅图像、38 个标注框，表中结果更适合作为阶段性验证结论，而不能直接视为大规模泛化性能结论。即便如此，mAP@0.5 与最终轮次的 Precision / Recall 已经表明，现有检测模型能够较稳定地输出海参目标位置，为后续视觉误差计算提供基础。":
        "需要说明的是，由于当前验证集规模仅为 12 幅图像、38 个标注框，表中结果更适合作为阶段性验证结论，而不能直接视为大规模泛化性能结论。即便如此，mAP@0.5 以及最终轮次的精确率和召回率已经表明，现有检测模型能够较稳定地输出海参目标位置，为后续视觉误差计算提供基础。",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            paragraph.text = replacements[text]


def set_table_borders(table) -> None:
    """为表格补齐外边框和内部网格线。"""

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_element = table._element
    table_pr = table_element.tblPr
    borders = table_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def polish_tables(doc: Document) -> None:
    """统一表格对齐、字体和边框。"""

    for table in doc.tables:
        set_table_borders(table)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.line_spacing = 1.15
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    set_paragraph_fonts(
                        paragraph,
                        chinese_font="宋体",
                        ascii_font="Times New Roman",
                        size_pt=10.5,
                        bold=(row_index == 0),
                    )


def polish_paragraph_layout(doc: Document) -> None:
    """统一段落的对齐、颜色、字体和标题样式。"""

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        # 图片段落不动，避免影响内嵌图片。
        if paragraph._element.xpath(".//pic:pic"):
            continue

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        if index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.2
            set_paragraph_fonts(paragraph, chinese_font="黑体", ascii_font="Times New Roman", size_pt=16, bold=True)
            continue

        if text == "摘要":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.25
            set_paragraph_fonts(paragraph, chinese_font="黑体", ascii_font="Times New Roman", size_pt=14, bold=True)
            continue

        if text.startswith("关键词："):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.5
            set_paragraph_fonts(paragraph, chinese_font="宋体", ascii_font="Times New Roman", size_pt=12, bold=False)
            continue

        if is_chapter_heading(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.25
            set_paragraph_fonts(paragraph, chinese_font="黑体", ascii_font="Times New Roman", size_pt=14, bold=True)
            continue

        if is_section_heading(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.25
            set_paragraph_fonts(paragraph, chinese_font="黑体", ascii_font="Times New Roman", size_pt=12, bold=True)
            continue

        if is_caption(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.2
            set_paragraph_fonts(paragraph, chinese_font="宋体", ascii_font="Times New Roman", size_pt=10.5, bold=False)
            continue

        if text.startswith("【此处插入"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.25
            set_paragraph_fonts(paragraph, chinese_font="楷体", ascii_font="Times New Roman", size_pt=11, bold=False)
            continue

        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = 1.5
        set_paragraph_fonts(paragraph, chinese_font="宋体", ascii_font="Times New Roman", size_pt=12, bold=False)


def main() -> None:
    """执行文档润色。"""

    doc = Document(SOURCE_DOC)
    replace_formula_like_paragraphs(doc)
    polish_paragraph_layout(doc)
    polish_tables(doc)
    doc.save(OUTPUT_DOC)
    print(f"generated: {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
